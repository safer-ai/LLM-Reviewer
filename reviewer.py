#!/usr/bin/env python3
"""
LLM Document Reviewer

A tool that uses LLM APIs to review and improve academic documents.
Supports both PDF and Word documents as input.
Implements a two-pass review:
1. Detailed review of text chunks.
2. Global review of the entire document for consistency and structure.
"""

import anthropic
import base64
import httpx
from typing import List, Dict, Union, Optional, Tuple, Any, Literal
import os
import time
from pathlib import Path
import argparse
import json
import re
from urllib.parse import urlparse
import logging
import yaml
import sys
from dataclasses import dataclass
import io

# Attempt to import fitz (PyMuPDF)
try:
    import fitz  # PyMuPDF
except ImportError:
    fitz = None

# Setup logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger('llm_reviewer')


class DocumentProcessor:
    """Handles loading documents and extracting text and structural elements."""

    def load_document(self, source: str) -> Tuple[bytes, str]:
        """
        Load a document from a URL or local path.

        Args:
            source: URL or local path to the document

        Returns:
            Tuple of (document_data_bytes, document_type_string)

        Raises:
            FileNotFoundError: If the local file doesn't exist
            ValueError: If the URL is invalid, inaccessible, or file type unsupported
            IOError: If there are issues reading the file
        """
        file_path = Path(source)
        file_ext = file_path.suffix.lower()

        parsed_url = urlparse(source)
        if parsed_url.scheme in ('http', 'https'):
            logger.info(f"Downloading document from URL: {source}")
            try:
                response = httpx.get(source, timeout=30)
                response.raise_for_status()
                content = response.content

                content_type = response.headers.get('Content-Type', '').lower()
                if 'pdf' in content_type:
                    return content, 'pdf'
                elif 'vnd.openxmlformats-officedocument.wordprocessingml.document' in content_type or file_ext == '.docx':
                    return content, 'docx'
                elif file_ext == '.pdf': # Fallback for URLs without proper content type
                    return content, 'pdf'
                else:
                    raise ValueError(f"Unsupported document type from URL. Content-Type: {content_type}, File extension: {file_ext}")

            except httpx.HTTPStatusError as e:
                raise ValueError(f"HTTP error when accessing {source}: {e.response.status_code}")
            except httpx.RequestError as e:
                raise ValueError(f"Network error when accessing {source}: {str(e)}")
        else:
            if not file_path.exists(): # This path is already resolved by ConfigManager
                raise FileNotFoundError(f"Document file not found: {file_path}") # No .resolve() needed here

            if file_ext == '.pdf':
                doc_data, doc_type_str = self._load_pdf_from_path(file_path)
            elif file_ext == '.docx':
                doc_data, doc_type_str = self._load_docx_from_path(file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_ext}. Supported types are .pdf and .docx")
            
            return doc_data, doc_type_str

    def _load_pdf_from_path(self, file_path: Path) -> Tuple[bytes, str]:
        logger.info(f"Loading PDF from local path: {file_path}")
        try:
            with open(file_path, 'rb') as f:
                content = f.read()
            if not content.startswith(b'%PDF'):
                logger.warning(f"File {file_path} may not be a valid PDF (missing %PDF header).")
            return content, 'pdf'
        except IOError as e:
            raise IOError(f"Error reading PDF file {file_path}: {str(e)}")

    def _load_docx_from_path(self, file_path: Path) -> Tuple[bytes, str]:
        logger.info(f"Loading DOCX from local path: {file_path}")
        try:
            with open(file_path, 'rb') as f:
                content = f.read()
            if not content.startswith(b'PK'): # DOCX are zip files, typically starting with 'PK'
                logger.warning(f"File {file_path} may not be a valid DOCX file (missing PK header).")
            return content, 'docx'
        except IOError as e:
            raise IOError(f"Error reading DOCX file {file_path}: {str(e)}")

    def get_full_text(self, doc_data: bytes, doc_type: str) -> str:
        """
        Extracts all text content from the document.

        Args:
            doc_data: Raw binary data of the document.
            doc_type: Type of the document ('pdf' or 'docx').

        Returns:
            The extracted text as a single string.

        Raises:
            ValueError: If the document type is unsupported or if required libraries are missing.
            RuntimeError: If text extraction fails.
        """
        logger.info(f"Extracting full text from {doc_type.upper()} document.")
        if doc_type == 'pdf':
            if fitz is None:
                raise ValueError("PyMuPDF (fitz) is required for PDF text extraction. Install with 'pip install PyMuPDF'")
            try:
                pdf_document = fitz.open(stream=doc_data, filetype="pdf")
                text_parts = []
                for page_num in range(len(pdf_document)):
                    page = pdf_document.load_page(page_num)
                    text_parts.append(page.get_text("text")) # "text" gives plain text
                full_text = "\n\n".join(text_parts) # Join pages with double newlines to preserve paragraph breaks better
                logger.info(f"Extracted {len(full_text)} characters from PDF ({len(pdf_document)} pages).")
                return full_text
            except Exception as e:
                raise RuntimeError(f"Failed to extract text from PDF: {str(e)}")
        elif doc_type == 'docx':
            try:
                full_text = self._extract_text_from_docx_data(doc_data)
                logger.info(f"Extracted {len(full_text)} characters from DOCX.")
                return full_text
            except Exception as e:
                raise RuntimeError(f"Failed to extract text from DOCX: {str(e)}")
        else:
            raise ValueError(f"Unsupported document type for text extraction: {doc_type}")

    def _extract_text_from_docx_data(self, docx_data: bytes) -> str:
        try:
            import docx # type: ignore
        except ImportError:
            raise ValueError("python-docx is required for DOCX processing. Install with 'pip install python-docx'")

        try:
            docx_file = io.BytesIO(docx_data)
            doc = docx.Document(docx_file)
            
            text_content = []
            # Paragraphs first
            for para in doc.paragraphs:
                text_content.append(para.text)
            
            # Then tables, trying to maintain some structure
            for table_idx, table in enumerate(doc.tables):
                text_content.append(f"\n--- Table {table_idx+1} Start ---")
                for row in table.rows:
                    row_texts = []
                    for cell in row.cells:
                        cell_text_parts = [p.text for p in cell.paragraphs]
                        row_texts.append(" | ".join(cell_text_parts)) # Join cell paragraphs, then join cells
                    text_content.append(" || ".join(row_texts)) # Double pipe for rows
                text_content.append(f"--- Table {table_idx+1} End ---\n")
            
            return "\n\n".join(text_content) # Join all parts with double newline
        except Exception as e:
            raise ValueError(f"Failed to parse DOCX data for text extraction: {str(e)}")


    def extract_structural_elements(self, full_text: str, doc_type: str, doc_data: Optional[bytes] = None) -> Dict[str, List[str]]:
        logger.info(f"Extracting structural elements from {doc_type.upper()} document.")
        elements: Dict[str, List[str]] = {
            'headings': [],
            'figure_captions': [],
            'table_captions': []
        }

        # DOCX: Try to use styles for headings first
        if doc_type == 'docx' and doc_data:
            try:
                import docx as docx_lib # type: ignore
                docx_file = io.BytesIO(doc_data)
                doc = docx_lib.Document(docx_file)
                for para in doc.paragraphs:
                    # Check for standard heading styles (Heading 1, Heading 2, etc.)
                    if para.style and para.style.name.lower().startswith('heading'):
                        trimmed_text = para.text.strip()
                        if trimmed_text:
                            elements['headings'].append(trimmed_text)
            except Exception as e:
                logger.warning(f"Could not use python-docx for DOCX heading extraction, falling back to regex: {e}")
                # Fall through to regex if docx parsing fails or no styled headings

        # PDF: (Future) Use fitz for more advanced structure detection (font sizes, ToC) if full_text is not enough
        # if doc_type == 'pdf' and doc_data and fitz:
        #     pdf_doc = fitz.open(stream=doc_data, filetype="pdf")
        #     toc = pdf_doc.get_toc() # Table of contents
        #     for level, title, page_num in toc:
        #         elements['headings'].append(title)
        # This is just an example, ToC extraction needs to be integrated carefully.

        # General Regex-based extraction (fallback or primary for PDFs for now)
        # Split text into lines for easier regex matching per line
        text_lines = full_text.splitlines()

        heading_patterns = [
            r"^\s*(?:[IVXLCDM]+\.|\d+(?:\.\d+)*)\s+([A-Z].{5,150})",  # Numbered headings (1. ..., 1.1 ..., IV. ...) - stricter on content
            r"^\s*(?:Chapter|Section|Part)\s+([IVXLCDM\d]+)\s*[:\-–—]?\s*(.{5,150})",    # Chapter X: Title, Section Y - Title
            r"^\s*(Abstract|Introduction|Conclusion|Summary|Recommendations|Acknowledgements|References|Appendix [A-Z0-9]+)\s*$", # Common section titles
            r"^\s*#{1,4}\s+(.{5,150})" # Markdown style headings
        ]
        # Only add if regex-based headings are needed (e.g. no DOCX style headings found)
        if not elements['headings'] or doc_type == 'pdf': # If headings list is empty or it's a PDF
            temp_headings = []
            for line in text_lines:
                line_stripped = line.strip()
                if not line_stripped or len(line_stripped) < 4 or len(line_stripped) > 200 : # Basic length filter
                    continue
                for pattern_idx, pattern in enumerate(heading_patterns):
                    match = re.match(pattern, line_stripped, re.IGNORECASE)
                    if match:
                        if pattern_idx == 1 and len(match.groups()) == 2: # For "Chapter X: Title"
                            heading_text = f"{match.group(1).strip()}: {match.group(2).strip()}"
                        elif match.groups():
                            heading_text = match.group(1).strip()
                        else: # Should not happen if pattern is correct
                            heading_text = line_stripped
                        
                        # Avoid adding very short or list-like items as headings
                        if len(heading_text) > 4 and not re.match(r"^\s*[\*\-\•\>]", heading_text):
                             temp_headings.append(heading_text)
                             break # Found a heading pattern for this line
            # De-duplicate while preserving order for regex-found headings
            seen_headings = set()
            for h in temp_headings:
                if h not in seen_headings:
                    elements['headings'].append(h)
                    seen_headings.add(h)


        figure_pattern = r"(?i)(?:Figure|Fig\.?|FIG\.?)\s+(\d+(?:[\.\-]\d+)?)\s*[:\-–—\.]?\s*(.+)"
        for line in text_lines:
            line_stripped = line.strip()
            match = re.search(figure_pattern, line_stripped) # search, not match, as it might not be at line start
            if match:
                num = match.group(1)
                caption_text = match.group(2).strip()
                if caption_text and len(caption_text) > 5: # Basic filter
                    elements['figure_captions'].append(f"Figure {num}: {caption_text}")
        
        table_pattern = r"(?i)(?:Table|Tab\.?|TAB\.?)\s+(\d+(?:[\.\-]\d+)?)\s*[:\-–—\.]?\s*(.+)"
        for line in text_lines:
            line_stripped = line.strip()
            match = re.search(table_pattern, line_stripped) # search, not match
            if match:
                num = match.group(1)
                caption_text = match.group(2).strip()
                if caption_text and len(caption_text) > 5: # Basic filter
                    elements['table_captions'].append(f"Table {num}: {caption_text}")
        
        logger.info(f"Found {len(elements['headings'])} potential headings.")
        logger.info(f"Found {len(elements['figure_captions'])} potential figure captions.")
        logger.info(f"Found {len(elements['table_captions'])} potential table captions.")
        return elements


class PromptLoader:
    """Handles loading and formatting prompt templates."""

    def __init__(self, prompt_paths: Dict[str, str]):
        self.prompt_paths = {name: Path(path) for name, path in prompt_paths.items() if isinstance(path, str)}
        self.prompts: Dict[str, str] = {}

        required_prompts_keys = ['pass1_system_prompt', 'pass1_review_prompt']
        
        for key, path_obj in self.prompt_paths.items():
            is_required = key in required_prompts_keys
            if not path_obj.exists():
                if is_required:
                    raise FileNotFoundError(f"Required prompt template not found: {key} at {path_obj}")
                else:
                    # For optional prompts like pass2, log if path specified but not found.
                    # If path wasn't specified for an optional prompt, it won't be in self.prompt_paths.
                    logger.warning(f"Optional prompt template not found: {key} at {path_obj}. This may be an issue if it's needed by config.")
                    continue # Don't try to load non-existent optional file
            
            try:
                self.prompts[key] = self._load_prompt(path_obj)
            except (IOError, ValueError) as e:
                if is_required: raise e # Re-raise if required prompt fails to load
                else: logger.error(f"Failed to load optional prompt {key} from {path_obj}: {e}")


        # Validate pass1_review_prompt placeholders
        if 'pass1_review_prompt' in self.prompts:
            pass1_rev_prompt = self.prompts['pass1_review_prompt']
            # {page_instructions_content} is optional placeholder for Pass 1
            missing_placeholders = [p for p in ['{instructions}', '{text_chunk}'] if p not in pass1_rev_prompt]
            if missing_placeholders:
                raise ValueError(f"Pass 1 review prompt template ({self.prompt_paths.get('pass1_review_prompt')}) is missing required placeholder(s): {', '.join(missing_placeholders)}")
        elif 'pass1_review_prompt' in self.prompt_paths: # Path was given, but file not loaded (e.g., empty)
             raise ValueError(f"Pass 1 review prompt template was configured but could not be loaded or is empty: {self.prompt_paths['pass1_review_prompt']}")


    def _load_prompt(self, path: Path) -> str:
        try:
            with open(path, 'r', encoding='utf-8') as f:
                prompt = f.read().strip()
            if not prompt:
                raise ValueError(f"Prompt template is empty: {path}")
            return prompt
        except IOError as e:
            raise IOError(f"Failed to load prompt template from {path}: {str(e)}")
        except ValueError as e: # Re-raise if empty
            raise e

    def get_prompt(self, prompt_name: str, is_required: bool = True) -> str:
        if prompt_name not in self.prompts:
            if is_required:
                # Path might exist in self.prompt_paths but failed to load into self.prompts
                path_for_prompt = self.prompt_paths.get(prompt_name, "Path not configured")
                raise FileNotFoundError(f"Required prompt '{prompt_name}' not loaded. Check configuration and file: {path_for_prompt}")
            else: # Optional prompt not found or not loaded
                logger.info(f"Optional prompt '{prompt_name}' not available.")
                return "" # Return empty string for optional prompts that are missing
        return self.prompts[prompt_name]

    def format_pass1_review_prompt(self, instructions: str, text_chunk: str,
                                   page_ranges: Optional[List[Tuple[int, int]]] = None) -> str:
        prompt_template = self.get_prompt('pass1_review_prompt') # is_required=True by default
        
        page_instructions_content = ""
        if page_ranges:
            page_instructions_content = "Please pay special attention to content that might correspond to these page ranges:\n"
            for start, end in page_ranges:
                page_instructions_content += f"- Page {start}" if start == end else f"- Pages {start} to {end}\n"
        
        formatted_prompt = prompt_template.replace('{instructions}', instructions)
        formatted_prompt = formatted_prompt.replace('{text_chunk}', text_chunk)
        
        # Optional placeholder for page instructions
        if '{page_instructions_content}' in formatted_prompt:
             formatted_prompt = formatted_prompt.replace('{page_instructions_content}', page_instructions_content)
        
        return formatted_prompt

    def format_pass2_global_review_prompt(self, full_document_text: str,
                                          structural_elements: Dict[str, List[str]]) -> str:
        prompt_template = self.get_prompt('pass2_review_prompt', is_required=True)
        
        formatted_elements_parts = []
        for key, items in structural_elements.items():
            if items:
                items_str = "\n".join([f"  - {item}" for item in items])
                formatted_elements_parts.append(f"{key.replace('_', ' ').title()}:\n{items_str}")
        structural_elements_text = "\n\n".join(formatted_elements_parts)
        if not structural_elements_text:
            structural_elements_text = "No structural elements were automatically extracted or provided."

        formatted_prompt = prompt_template.replace('{full_document_text}', full_document_text)
        formatted_prompt = formatted_prompt.replace('{structural_elements}', structural_elements_text)
        return formatted_prompt


@dataclass
class APIResponse:
    response: str
    model: str
    input_tokens: Optional[int] = None
    output_tokens: Optional[int] = None
    error: Optional[str] = None


class LLMClient:
    def __init__(self, api_key: Optional[str] = None, model: str = "claude-3-opus-20240229",
                 max_tokens: int = 4096, temperature: float = 0.0,
                 retry_attempts: int = 3, retry_delay: int = 5):
        self.api_key = api_key or os.environ.get("ANTHROPIC_API_KEY")
        if not self.api_key:
            raise ValueError("Anthropic API key is required. Set ANTHROPIC_API_KEY or configure.")
        self.client = anthropic.Anthropic(api_key=self.api_key)
        self.model = model
        self.max_tokens = max_tokens
        self.temperature = temperature
        self.retry_attempts = retry_attempts
        self.retry_delay = retry_delay

    def review_text_content(self, system_prompt: str, user_prompt: str) -> APIResponse:
        last_error = None
        for attempt in range(self.retry_attempts):
            try:
                logger.info(f"API call (attempt {attempt+1}/{self.retry_attempts}) model: {self.model}, system prompt: {len(system_prompt)} chars, user_prompt: {len(user_prompt)} chars.")
                if attempt > 0:
                    sleep_time = self.retry_delay * (2 ** (attempt - 1))
                    logger.info(f"Retrying in {sleep_time} seconds...")
                    time.sleep(sleep_time)
                
                response_obj = self.client.messages.create(
                    model=self.model,
                    max_tokens=self.max_tokens,
                    temperature=self.temperature,
                    system=system_prompt,
                    messages=[{"role": "user", "content": user_prompt}],
                )
                
                input_tokens = getattr(response_obj.usage, 'input_tokens', None)
                output_tokens = getattr(response_obj.usage, 'output_tokens', None)
                
                if not response_obj.content or not isinstance(response_obj.content, list) or not response_obj.content[0].text:
                    raise anthropic.APIError("Malformed API response: Empty or unexpected content structure.")

                return APIResponse(
                    response=response_obj.content[0].text, model=self.model,
                    input_tokens=input_tokens, output_tokens=output_tokens
                )
            except anthropic.RateLimitError as e: last_error = e; logger.warning(f"Rate limit (attempt {attempt+1}): {e}"); time.sleep(self.retry_delay * 3)
            except anthropic.APITimeoutError as e: last_error = e; logger.warning(f"Timeout (attempt {attempt+1}): {e}")
            except anthropic.APIConnectionError as e: last_error = e; logger.warning(f"Connection error (attempt {attempt+1}): {e}")
            except anthropic.APIStatusError as e: # More specific for HTTP status errors
                last_error = e
                logger.warning(f"API Status Error (attempt {attempt+1}): {e.status_code} - {e.message}")
                if e.status_code == 400: # Bad request, often due to prompt issues or model limits
                    logger.error(f"Bad Request (400) from API. This might be due to prompt length or content. Details: {e.response.text}")
                    # No point retrying a 400 error that's due to bad input.
                    return APIResponse(response="", model=self.model, error=f"API Bad Request (400): {e.message}. Details: {e.response.text if e.response else 'N/A'}")
                elif e.status_code == 429: # Explicit rate limit, caught by RateLimitError too
                     time.sleep(self.retry_delay * 5) # Longer sleep for explicit 429
            except anthropic.APIError as e: last_error = e; logger.warning(f"Other Anthropic API error (attempt {attempt+1}): {e}")
            except Exception as e: last_error = e; logger.error(f"Unexpected error during API call (attempt {attempt+1}): {e}", exc_info=True)
        
        error_message = f"All {self.retry_attempts} API call attempts failed for model {self.model}. Last error: {str(last_error)}"
        logger.error(error_message)
        return APIResponse(response="", model=self.model, error=error_message)


class ReviewFormatter:
    def __init__(self):
        self.suggestion_pattern = re.compile(
            r'\(change\s+"((?:[^"\\]|\\.|"(?:\\.|[^"\\])*")*?)"\s+->\s+"((?:[^"\\]|\\.|"(?:\\.|[^"\\])*")*?)"\)',
            re.IGNORECASE | re.DOTALL
        )

    def parse_suggestions(self, response: str) -> List[Dict[str, str]]:
        if not response: return []
        suggestions = []
        matches = self.suggestion_pattern.findall(response)
        for original, improved in matches:
            original_clean = original.strip().replace('\\"', '"')
            improved_clean = improved.strip().replace('\\"', '"')
            if original_clean and improved_clean and original_clean != improved_clean:
                suggestions.append({'original': original_clean, 'improved': improved_clean})
        logger.info(f"Parsed {len(suggestions)} suggestions from response.")
        return suggestions


class ConfigManager:
    def __init__(self, config_path: str):
        self.config_path = Path(config_path) # Path is already resolved in main()
        if not self.config_path.exists():
            raise FileNotFoundError(f"Configuration file not found: {self.config_path}")
        
        self.config = self._load_config()
        self._validate_config()

    def _load_config(self) -> Dict[str, Any]:
        try:
            with open(self.config_path, 'r', encoding='utf-8') as f: config = yaml.safe_load(f)
            logger.info(f"Configuration loaded from {self.config_path}")
            return config or {}
        except yaml.YAMLError as e: raise ValueError(f"Invalid YAML in {self.config_path}: {e}")
        except IOError as e: raise IOError(f"Failed to load config from {self.config_path}: {e}")

    def _validate_config(self):
        required_sections = ['api', 'document', 'output', 'prompts', 'review']
        for section in required_sections:
            if section not in self.config: raise ValueError(f"Missing section: '{section}'")

        if not self.get_document_config().get('path'): raise ValueError("'document.path' is required.")

        prompt_cfg = self.get_prompt_config_raw() # Use raw for validation before resolving
        required_prompts = ['pass1_system_prompt', 'pass1_review_prompt']
        for p_key in required_prompts:
            if not prompt_cfg.get(p_key): raise ValueError(f"'prompts.{p_key}' is required.")

        if self.get_enable_global_review():
            for p_key in ['pass2_system_prompt', 'pass2_review_prompt']:
                if not prompt_cfg.get(p_key): raise ValueError(f"Global review enabled, missing 'prompts.{p_key}'.")
        
        review_cfg = self.get_review_config()
        for key, type_check, condition in [
            ('chunk_size_words', lambda x: isinstance(x, int), lambda x: x > 0),
            ('chunk_overlap_words', lambda x: isinstance(x, int), lambda x: x >= 0)
        ]:
            val = review_cfg.get(key)
            if val is None or not type_check(val) or not condition(val):
                raise ValueError(f"'review.{key}' must be a {'positive' if key=='chunk_size_words' else 'non-negative'} integer. Found: {val}")
        if review_cfg.get('chunk_overlap_words', 0) >= review_cfg.get('chunk_size_words', 1):
             raise ValueError("'review.chunk_overlap_words' must be less than 'review.chunk_size_words'.")

    def get_api_config(self) -> Dict[str, Any]: return self.config.get('api', {})
    def get_document_config(self) -> Dict[str, Any]: return self.config.get('document', {})
    def get_output_config(self) -> Dict[str, Any]: return self.config.get('output', {})
    def get_review_config(self) -> Dict[str, Any]: return self.config.get('review', {})

    def get_prompt_config_raw(self) -> Dict[str, Any]: # Helper for validation
        return self.config.get('prompts', {})

    def get_prompt_config(self) -> Dict[str, str]:
        raw_paths = self.config.get('prompts', {})
        if not isinstance(raw_paths, dict): return {}
        resolved = {}
        base_dir = self.config_path.parent
        for name, rel_path in raw_paths.items():
            if isinstance(rel_path, str):
                p = Path(rel_path)
                resolved[name] = str(p.resolve() if p.is_absolute() else (base_dir / p).resolve())
        return resolved

    def get_enable_global_review(self) -> bool: return self.get_review_config().get('enable_global_review', False)
    def get_chunk_size_words(self) -> int: return self.get_review_config().get('chunk_size_words', 700)
    def get_chunk_overlap_words(self) -> int: return self.get_review_config().get('chunk_overlap_words', 100)

    def get_resolved_document_path(self) -> str:
        path_str = self.get_document_config().get('path')
        if not path_str: raise ValueError("'document.path' is missing.")
        parsed = urlparse(path_str)
        if parsed.scheme in ('http', 'https'): return path_str
        p = Path(path_str)
        return str(p.resolve() if p.is_absolute() else (self.config_path.parent / p).resolve())

    def get_resolved_output_path(self) -> Optional[str]:
        path_str = self.get_output_config().get('path')
        if not path_str: return None
        p = Path(path_str)
        return str(p.resolve() if p.is_absolute() else (self.config_path.parent / p).resolve())

    def parse_page_ranges(self) -> Optional[List[Tuple[int, int]]]:
        pages_str = self.get_document_config().get('pages')
        if not pages_str: return None
        ranges = []
        try:
            for r_part in str(pages_str).split(','):
                r = r_part.strip()
                if '-' in r:
                    s, e = map(int, r.split('-', 1))
                    if not (1 <= s <= e): raise ValueError(f"Invalid range {r}")
                    ranges.append((s, e))
                else:
                    p = int(r); 
                    if p < 1: raise ValueError(f"Invalid page {p}")
                    ranges.append((p, p))
            logger.info(f"Parsed page ranges: {ranges}")
            return ranges
        except ValueError as e: raise ValueError(f"Invalid page range format '{pages_str}': {e}")


class DocumentReviewer:
    def __init__(self, config_manager: ConfigManager):
        self.config_manager = config_manager
        self._initialize_components()

    def _initialize_components(self):
        api_cfg = self.config_manager.get_api_config()
        prompt_paths = self.config_manager.get_prompt_config()

        self.document_processor = DocumentProcessor()
        self.prompt_loader = PromptLoader(prompt_paths=prompt_paths)
        self.llm_client = LLMClient(
            api_key=api_cfg.get('key'), model=api_cfg.get('model', "claude-3-sonnet-20240229"),
            max_tokens=api_cfg.get('max_tokens', 4096), temperature=api_cfg.get('temperature', 0.0),
            retry_attempts=api_cfg.get('retry_attempts', 3), retry_delay=api_cfg.get('retry_delay', 5)
        )
        self.review_formatter = ReviewFormatter()

    def _chunk_text(self, text: str, chunk_size_words: int, overlap_words: int) -> List[str]:
        if not isinstance(text, str) or not text.strip():
            logger.warning("Input text for chunking is empty or not a string. Returning no chunks.")
            return []
            
        if not isinstance(chunk_size_words, int) or chunk_size_words <= 0:
            logger.warning(f"Chunk size must be a positive integer, got {chunk_size_words}. Returning text as a single chunk.")
            return [text]
        if not isinstance(overlap_words, int) or overlap_words < 0:
            logger.warning(f"Overlap words must be a non-negative integer, got {overlap_words}. Setting to 0.")
            overlap_words = 0
        
        # Ensure overlap is reasonably less than chunk_size
        if overlap_words >= chunk_size_words:
            logger.warning(f"Overlap ({overlap_words}) >= chunk size ({chunk_size_words}). Adjusting overlap to {chunk_size_words // 4}.")
            overlap_words = max(0, chunk_size_words // 4) # Ensure overlap is not negative if chunk_size is small

        words = text.split() # Simple split by any whitespace, creates a list of actual words
        if not words:
            logger.warning("Text resulted in an empty list of words after splitting. Returning no chunks.")
            return []

        chunks: List[str] = []
        current_word_idx = 0
        text_word_count = len(words)

        while current_word_idx < text_word_count:
            # Determine the end of the current chunk
            end_word_idx = min(current_word_idx + chunk_size_words, text_word_count)
            
            # Slice the words list to get the current chunk's words
            current_chunk_word_list = words[current_word_idx:end_word_idx]
            
            if not current_chunk_word_list: # Should not happen if loop condition is correct
                break 
                
            chunks.append(" ".join(current_chunk_word_list)) # Join the words with single spaces

            # Determine the starting point for the next chunk
            # Step must ensure progress and be at least 1 word.
            step = chunk_size_words - overlap_words
            if step <= 0:
                # If chunk_size is small or overlap is too large, ensure a minimum step.
                # For example, if chunk_size is 10 and overlap is 9, step is 1.
                # If chunk_size is 10 and overlap is 10, step would be 0, so fix it.
                step = max(1, chunk_size_words // 2 if chunk_size_words > 1 else 1)
            
            current_word_idx += step

            # Safety break if somehow stuck, though step calculation should prevent this.
            if step == 0 and current_word_idx < text_word_count : # Should ideally not be needed
                 logger.error("Chunking step became 0, forcing break to prevent infinite loop.")
                 break
                 
        if not chunks and words: # If original text had words but no chunks were made (e.g. logic error)
            logger.warning("Chunking resulted in zero chunks from non-empty text. Returning text as a single chunk as a fallback.")
            return [text] # Fallback

        logger.info(f"Split text into {len(chunks)} chunks (target size ~{chunk_size_words} words, overlap ~{overlap_words} words).")
        return chunks

    def review_document(self) -> Dict[str, Any]:
        doc_source = self.config_manager.get_resolved_document_path()
        logger.info(f"Starting review of document: {doc_source}")

        results: Dict[str, Any] = {
            'document_source': doc_source,
            'document_type': None,
            'pass1_detailed_review': {
                'status': 'Not Run', 'total_suggestions': 0, 'suggestions': [],
                'total_input_tokens': 0, 'total_output_tokens': 0, 'chunks_processed': 0
            },
            'pass2_global_review': {'status': 'Not Run', 'report': None, 'input_tokens': 0, 'output_tokens': 0},
            'model_info': {'pass1_model': None, 'pass2_model': None},
            'errors': []
        }

        try:
            doc_data, doc_type = self.document_processor.load_document(doc_source)
            results['document_type'] = doc_type
            logger.info(f"Successfully loaded {doc_type.upper()} document ({len(doc_data)/1024:.1f} KB)")

            full_text = self.document_processor.get_full_text(doc_data, doc_type)
            if not full_text.strip():
                logger.warning("Extracted text is empty or whitespace only. Cannot proceed with review.")
                results['errors'].append("Extracted text is empty.")
                self._save_results(results); return results # Early exit

            structural_elements = self.document_processor.extract_structural_elements(full_text, doc_type, doc_data)

        except Exception as e:
            logger.error(f"Failed during document loading or initial processing: {str(e)}", exc_info=True)
            results['errors'].append(f"Initial processing error: {str(e)}")
            self._save_results(results); return results

        # --- Pass 1: Detailed Chunk Review ---
        logger.info("--- Starting Pass 1: Detailed Chunk Review ---")
        pass1_results = results['pass1_detailed_review']
        try:
            pass1_sys_prompt = self.prompt_loader.get_prompt('pass1_system_prompt')
            review_instructions = self.config_manager.get_review_config().get('instructions', '')
            page_ranges_for_focus = self.config_manager.parse_page_ranges() # Optional

            chunk_size = self.config_manager.get_chunk_size_words()
            overlap = self.config_manager.get_chunk_overlap_words()
            text_chunks = self._chunk_text(full_text, chunk_size, overlap)
            
            if not text_chunks:
                logger.warning("No text chunks generated for Pass 1 review.")
                pass1_results['status'] = 'Skipped - No text chunks'
            else:
                pass1_results['status'] = 'In Progress'
                all_pass1_suggestions = []
                for i, chunk in enumerate(text_chunks):
                    logger.info(f"Processing Pass 1 - Chunk {i+1}/{len(text_chunks)} ({len(chunk)} chars)")
                    pass1_user_prompt = self.prompt_loader.format_pass1_review_prompt(
                        instructions=review_instructions, text_chunk=chunk, page_ranges=page_ranges_for_focus
                    )
                    api_response = self.llm_client.review_text_content(pass1_sys_prompt, pass1_user_prompt)

                    if api_response.error:
                        logger.error(f"API error in Pass 1, Chunk {i+1}: {api_response.error}")
                        results['errors'].append(f"Pass 1, Chunk {i+1} API Error: {api_response.error}")
                        # Decide if we should continue to next chunk or stop
                        continue # Continue with next chunk for now

                    chunk_suggestions = self.review_formatter.parse_suggestions(api_response.response)
                    for sug_idx, sug in enumerate(chunk_suggestions):
                        sug['id'] = f"p1_c{i+1}_s{sug_idx+1}" # Unique ID for suggestion
                        sug['chunk_index'] = i+1
                    all_pass1_suggestions.extend(chunk_suggestions)
                    
                    pass1_results['total_input_tokens'] += api_response.input_tokens or 0
                    pass1_results['total_output_tokens'] += api_response.output_tokens or 0
                    pass1_results['chunks_processed'] += 1
                
                pass1_results['suggestions'] = all_pass1_suggestions
                pass1_results['total_suggestions'] = len(all_pass1_suggestions)
                pass1_results['status'] = 'Completed'
            results['model_info']['pass1_model'] = self.llm_client.model

        except Exception as e:
            logger.error(f"Error during Pass 1 review: {str(e)}", exc_info=True)
            results['errors'].append(f"Pass 1 Error: {str(e)}")
            pass1_results['status'] = 'Failed'

        # --- Pass 2: Global Document Review ---
        if self.config_manager.get_enable_global_review():
            logger.info("--- Starting Pass 2: Global Document Review ---")
            pass2_results = results['pass2_global_review']
            try:
                pass2_sys_prompt = self.prompt_loader.get_prompt('pass2_system_prompt', is_required=True)
                pass2_user_prompt = self.prompt_loader.format_pass2_global_review_prompt(full_text, structural_elements)
                
                pass2_results['status'] = 'In Progress'
                api_response = self.llm_client.review_text_content(pass2_sys_prompt, pass2_user_prompt)

                if api_response.error:
                    logger.error(f"API error in Pass 2 (Global Review): {api_response.error}")
                    results['errors'].append(f"Pass 2 API Error: {api_response.error}")
                    pass2_results['status'] = 'Failed'
                else:
                    pass2_results['report'] = api_response.response
                    pass2_results['input_tokens'] = api_response.input_tokens or 0
                    pass2_results['output_tokens'] = api_response.output_tokens or 0
                    pass2_results['status'] = 'Completed'
                results['model_info']['pass2_model'] = self.llm_client.model

            except FileNotFoundError as e: # Catch missing pass2 prompts specifically
                logger.error(f"Pass 2 prompt file not found, skipping global review: {e}")
                results['errors'].append(f"Pass 2 skipped (prompt missing): {e}")
                pass2_results['status'] = 'Skipped - Prompt Missing'
            except Exception as e:
                logger.error(f"Error during Pass 2 review: {str(e)}", exc_info=True)
                results['errors'].append(f"Pass 2 Error: {str(e)}")
                pass2_results['status'] = 'Failed'
        else:
            results['pass2_global_review']['status'] = 'Disabled by configuration'

        self._save_results(results)
        return results

    def _extract_summary(self, response: str) -> str: # General utility, keep as is
        if not response: return ""
        summary_patterns = [
            r"(?i)\b(?:summary|overall assessment|key findings|conclusion|final thoughts)\b\s*[:\-–—]?\s*\n*(.*?)(?=\n\n\S|\Z)",
            r"(?i)(summary|key issues|overall|conclusion):(.*?)(?=\n\n|\Z)"
        ] # Ensure re.DOTALL is used if these patterns cross newlines significantly. The first one is more robust.
        for pattern in summary_patterns:
            match = re.search(pattern, response, re.DOTALL | re.IGNORECASE)
            if match:
                summary_text = match.group(1) or (match.group(2) if len(match.groups()) > 1 else None)
                if summary_text and summary_text.strip(): return summary_text.strip()
        
        paragraphs = [p.strip() for p in response.split('\n\n') if p.strip()]
        if paragraphs:
            last_para = paragraphs[-1]
            if 30 < len(last_para) < 1000 and any(w in last_para.lower() for w in ['overall', 'summary', 'conclusion', 'general', 'finally', 'in sum']):
                return last_para
        return ""

    def _save_results(self, results: Dict[str, Any]):
        output_path_str = self.config_manager.get_resolved_output_path()
        if not output_path_str:
            logger.info("No output path configured. Results will not be saved to file.")
            return
            
        try:
            output_file = Path(output_path_str)
            output_file.parent.mkdir(parents=True, exist_ok=True)
            with open(output_file, 'w', encoding='utf-8') as f:
                json.dump(results, f, indent=2, ensure_ascii=False)
            logger.info(f"Results saved to {output_file.resolve()}")
        except IOError as e:
            logger.error(f"Failed to save results to {output_path_str}: {str(e)}")


def main():
    parser = argparse.ArgumentParser(
        description='LLM Document Reviewer - Review academic papers using Claude API',
        epilog='Example: reviewer.py --config parameters.yaml'
    )
    parser.add_argument('--config', '-c', default='parameters.yaml',
                       help='Path to the YAML configuration file (default: parameters.yaml)')
    parser.add_argument('--quiet', '-q', action='store_true', help='Suppress informational output')
    parser.add_argument('--version', '-v', action='version', version='LLM Document Reviewer v2.0.0') # Version bump
    
    args = parser.parse_args()
    
    if args.quiet:
        logging.getLogger().setLevel(logging.WARNING)
        logging.getLogger('llm_reviewer').setLevel(logging.WARNING)

    try:
        config_file_to_load = args.config
        if args.config == parser.get_default('config'):
            script_dir = Path(__file__).parent.resolve()
            default_config_near_script = script_dir / args.config
            if default_config_near_script.exists():
                config_file_to_load = str(default_config_near_script)
        
        config_manager = ConfigManager(config_file_to_load)
        reviewer = DocumentReviewer(config_manager)
        
        start_time = time.time()
        results = reviewer.review_document()
        end_time = time.time()
        
        output_path = config_manager.get_resolved_output_path()
        if not output_path and not results.get('errors'):
            print(json.dumps(results, indent=2, ensure_ascii=False))
        
        if results.get('errors'):
            logger.error(f"Review process completed with errors: {results.get('errors')}")
            print(f"Review completed with errors. Check logs and results file. First error: {results.get('errors')[0]}", file=sys.stderr)
            # Continue to print summary despite errors
        
        print(f"\n--- Review Summary ---")
        print(f"Processed document: {results.get('document_source', 'N/A')}")
        print(f"Total time: {end_time - start_time:.1f} seconds.")

        pass1_summary = results.get('pass1_detailed_review', {})
        print(f"\nPass 1 (Detailed Chunk Review):")
        print(f"  Status: {pass1_summary.get('status', 'N/A')}")
        print(f"  Chunks processed: {pass1_summary.get('chunks_processed', 0)}")
        print(f"  Suggestions found: {pass1_summary.get('total_suggestions', 0)}")
        print(f"  Input tokens: {pass1_summary.get('total_input_tokens', 0)}")
        print(f"  Output tokens: {pass1_summary.get('total_output_tokens', 0)}")

        pass2_summary = results.get('pass2_global_review', {})
        print(f"\nPass 2 (Global Document Review):")
        print(f"  Status: {pass2_summary.get('status', 'N/A')}")
        if pass2_summary.get('status') == 'Completed':
            print(f"  Report generated (see results file).")
            print(f"  Input tokens: {pass2_summary.get('input_tokens', 0)}")
            print(f"  Output tokens: {pass2_summary.get('output_tokens', 0)}")

        if output_path:
            print(f"\nFull results saved to: {output_path}")
        else:
            print("\nFull results printed above (or configure output.path to save to file).")

        if results.get('errors'):
            sys.exit(1) # Exit with error code if there were errors during processing
            
    except (FileNotFoundError, ValueError, IOError) as e:
        logger.error(f"Setup Error: {str(e)}", exc_info=not args.quiet)
        print(f"Error: {str(e)}", file=sys.stderr)
        sys.exit(1)
    except Exception as e:
        logger.exception("An unexpected error occurred:")
        print(f"Unexpected error: {str(e)}", file=sys.stderr)
        sys.exit(1)

if __name__ == '__main__':
    main()