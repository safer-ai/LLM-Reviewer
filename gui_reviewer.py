#!/usr/bin/env python3
"""
LLM Document Reviewer - GUI Version
A user-friendly interface for reviewing documents with AI

No command line needed - just double-click to run!
"""

import tkinter as tk
from tkinter import ttk, filedialog, messagebox, scrolledtext
import threading
import os
import sys
from pathlib import Path
import json
import webbrowser
from datetime import datetime
import queue

# Try to import the reviewer modules, with user-friendly error messages
try:
    from reviewer import (
        DocumentProcessor, PromptLoader, LLMClient, 
        ReviewFormatter, ConfigManager, DocumentReviewer
    )
except ImportError as e:
    import tkinter.messagebox as mb
    mb.showerror("Missing Files", 
                 "Required reviewer files are missing. Please ensure all files are in the same folder.")
    sys.exit(1)


class ReviewerGUI:
    def __init__(self, root):
        self.root = root
        self.root.title("Document Reviewer - AI-Powered Review Assistant")
        self.root.geometry("900x700")
        
        # Set icon if available
        try:
            if sys.platform.startswith('win'):
                self.root.iconbitmap('icon.ico')
        except:
            pass
        
        # Variables
        self.selected_file = tk.StringVar()
        self.status_var = tk.StringVar(value="Ready to review")
        self.progress_var = tk.DoubleVar()
        self.review_thread = None
        self.suggestions = []
        self.config_path = self._find_config_file()
        
        # Message queue for thread communication
        self.message_queue = queue.Queue()
        
        # Create GUI
        self._create_widgets()
        self._center_window()
        
        # Check for API key
        self._check_api_key()
        
        # Start message queue processor
        self.root.after(100, self._process_queue)
    
    def _find_config_file(self):
        """Find the configuration file"""
        # Look for config in same directory as script
        script_dir = Path(__file__).parent
        config_file = script_dir / "parameters.yaml"
        
        if not config_file.exists():
            # Create a default config
            self._create_default_config(config_file)
        
        return str(config_file)
    
    def _create_default_config(self, config_path):
        """Create a default configuration file"""
        default_config = """# LLM Document Reviewer Configuration

# API Configuration
api:
  # API key will be read from ANTHROPIC_API_KEY environment variable
  key: null
  model: "claude-3-sonnet-20241022"
  max_tokens: 4096
  temperature: 0.3
  retry_attempts: 3
  retry_delay: 2

# Document Configuration
document:
  path: null  # Will be set by GUI
  pages: null

# Output Configuration
output:
  path: null  # Will be set by GUI

# Prompt Configuration
prompts:
  pass1_system_prompt: "prompts/pass1_system_prompt.txt"
  pass1_review_prompt: "prompts/pass1_review_prompt.txt"
  pass2_system_prompt: "prompts/pass2_system_prompt.txt"
  pass2_review_prompt: "prompts/pass2_review_prompt.txt"

# Review Configuration
review:
  instructions: "Focus on clarity, grammar, and consistency."
  enable_global_review: true
  chunk_size_words: 700
  chunk_overlap_words: 100

# Batch Processing Configuration
batch:
  pages_per_batch: 3
"""
        
        # Also create prompts directory and files
        prompts_dir = config_path.parent / "prompts"
        prompts_dir.mkdir(exist_ok=True)
        
        # Create basic prompt files
        self._create_prompt_files(prompts_dir)
        
        with open(config_path, 'w') as f:
            f.write(default_config)
    
    def _create_prompt_files(self, prompts_dir):
        """Create default prompt files"""
        prompts = {
            "pass1_system_prompt.txt": """You are an expert academic editor reviewing documents.
Focus on grammar, spelling, clarity, and consistency.
Be thorough but constructive in your feedback.""",
            
            "pass1_review_prompt.txt": """Please review this text and suggest improvements.

Text to review:
{text_chunk}

Additional instructions: {instructions}

For each issue, format your response as:
(change "original text" -> "improved text")

{page_instructions}""",
            
            "pass2_system_prompt.txt": """You are an expert academic editor performing a final review.
Focus on document-wide consistency and structure.""",
            
            "pass2_review_prompt.txt": """Review the complete document for consistency and structure.

Document:
{full_document_text}

Structural elements:
{structural_elements}

Provide a comprehensive review report."""
        }
        
        for filename, content in prompts.items():
            filepath = prompts_dir / filename
            if not filepath.exists():
                with open(filepath, 'w') as f:
                    f.write(content)
    
    def _check_api_key(self):
        """Check if API key is configured"""
        if not os.environ.get("ANTHROPIC_API_KEY"):
            response = messagebox.askyesno(
                "API Key Required",
                "No Anthropic API key found.\n\n"
                "Would you like to set it now?\n\n"
                "You'll need your API key from:\n"
                "https://console.anthropic.com/api-keys"
            )
            
            if response:
                self._show_api_key_dialog()
    
    def _show_api_key_dialog(self):
        """Show dialog to enter API key"""
        dialog = tk.Toplevel(self.root)
        dialog.title("Enter API Key")
        dialog.geometry("500x200")
        dialog.transient(self.root)
        dialog.grab_set()
        
        # Center the dialog
        dialog.update_idletasks()
        x = (dialog.winfo_screenwidth() // 2) - (dialog.winfo_width() // 2)
        y = (dialog.winfo_screenheight() // 2) - (dialog.winfo_height() // 2)
        dialog.geometry(f"+{x}+{y}")
        
        # Instructions
        ttk.Label(dialog, text="Enter your Anthropic API Key:", 
                 font=("Arial", 10)).pack(pady=10)
        
        ttk.Label(dialog, text="Get your key from: https://console.anthropic.com/api-keys",
                 foreground="blue", cursor="hand2").pack()
        
        # API key entry
        key_var = tk.StringVar()
        entry = ttk.Entry(dialog, textvariable=key_var, width=60, show="*")
        entry.pack(pady=10, padx=20)
        
        def save_key():
            key = key_var.get().strip()
            if key:
                # Save to environment variable for this session
                os.environ["ANTHROPIC_API_KEY"] = key
                
                # Also create a batch file for Windows users
                if sys.platform.startswith('win'):
                    self._create_batch_file(key)
                
                messagebox.showinfo("Success", "API key saved for this session!")
                dialog.destroy()
            else:
                messagebox.showerror("Error", "Please enter a valid API key")
        
        ttk.Button(dialog, text="Save", command=save_key).pack(pady=10)
        
        entry.focus()
        entry.bind('<Return>', lambda e: save_key())
    
    def _create_batch_file(self, api_key):
        """Create a batch file that sets the API key and runs the program"""
        batch_content = f"""@echo off
set ANTHROPIC_API_KEY={api_key}
python "{Path(__file__).resolve()}"
pause
"""
        
        batch_file = Path(__file__).parent / "run_reviewer.bat"
        with open(batch_file, 'w') as f:
            f.write(batch_content)
    
    def _create_widgets(self):
        """Create all GUI widgets"""
        # Main container
        main_frame = ttk.Frame(self.root, padding="10")
        main_frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        
        # Configure grid weights
        self.root.columnconfigure(0, weight=1)
        self.root.rowconfigure(0, weight=1)
        main_frame.columnconfigure(1, weight=1)
        main_frame.rowconfigure(3, weight=1)
        
        # Title
        title_label = ttk.Label(main_frame, text="AI Document Reviewer", 
                               font=("Arial", 16, "bold"))
        title_label.grid(row=0, column=0, columnspan=3, pady=(0, 20))
        
        # File selection
        ttk.Label(main_frame, text="Select Document:").grid(row=1, column=0, 
                                                           sticky=tk.W, pady=5)
        
        file_entry = ttk.Entry(main_frame, textvariable=self.selected_file, 
                              state="readonly", width=50)
        file_entry.grid(row=1, column=1, sticky=(tk.W, tk.E), pady=5, padx=(10, 10))
        
        browse_btn = ttk.Button(main_frame, text="Browse...", 
                               command=self._select_file)
        browse_btn.grid(row=1, column=2, pady=5)
        
        # Review button
        self.review_btn = ttk.Button(main_frame, text="Start Review", 
                                    command=self._start_review,
                                    style="Accent.TButton")
        self.review_btn.grid(row=2, column=0, columnspan=3, pady=20)
        
        # Progress bar
        self.progress_bar = ttk.Progressbar(main_frame, variable=self.progress_var,
                                          maximum=100, length=400)
        self.progress_bar.grid(row=3, column=0, columnspan=3, sticky=(tk.W, tk.E),
                              pady=(0, 10))
        
        # Status label
        status_label = ttk.Label(main_frame, textvariable=self.status_var)
        status_label.grid(row=4, column=0, columnspan=3, pady=(0, 10))
        
        # Results area
        results_frame = ttk.LabelFrame(main_frame, text="Review Results", 
                                      padding="10")
        results_frame.grid(row=5, column=0, columnspan=3, sticky=(tk.W, tk.E, tk.N, tk.S),
                          pady=(10, 0))
        results_frame.columnconfigure(0, weight=1)
        results_frame.rowconfigure(0, weight=1)
        
        # Results text area
        self.results_text = scrolledtext.ScrolledText(results_frame, wrap=tk.WORD,
                                                     width=80, height=20)
        self.results_text.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        
        # Configure text tags for formatting
        self.results_text.tag_configure("heading", font=("Arial", 12, "bold"),
                                       foreground="darkblue")
        self.results_text.tag_configure("original", foreground="red")
        self.results_text.tag_configure("improved", foreground="green")
        self.results_text.tag_configure("category", font=("Arial", 10, "italic"))
        
        # Export buttons (initially hidden)
        self.export_frame = ttk.Frame(main_frame)
        
        ttk.Button(self.export_frame, text="Export to Word",
                  command=lambda: self._export_results("docx")).pack(side=tk.LEFT, padx=5)
        ttk.Button(self.export_frame, text="Export to Excel", 
                  command=lambda: self._export_results("xlsx")).pack(side=tk.LEFT, padx=5)
        ttk.Button(self.export_frame, text="Export to Markdown",
                  command=lambda: self._export_results("md")).pack(side=tk.LEFT, padx=5)
    
    def _center_window(self):
        """Center the window on screen"""
        self.root.update_idletasks()
        width = self.root.winfo_width()
        height = self.root.winfo_height()
        x = (self.root.winfo_screenwidth() // 2) - (width // 2)
        y = (self.root.winfo_screenheight() // 2) - (height // 2)
        self.root.geometry(f"{width}x{height}+{x}+{y}")
    
    def _select_file(self):
        """Open file dialog to select document"""
        filename = filedialog.askopenfilename(
            title="Select Document to Review",
            filetypes=[
                ("Supported Documents", "*.pdf;*.docx"),
                ("PDF Files", "*.pdf"),
                ("Word Documents", "*.docx"),
                ("All Files", "*.*")
            ]
        )
        
        if filename:
            self.selected_file.set(filename)
            self.results_text.delete(1.0, tk.END)
            self.export_frame.grid_forget()
    
    def _start_review(self):
        """Start the review process in a separate thread"""
        if not self.selected_file.get():
            messagebox.showwarning("No File Selected", 
                                 "Please select a document to review.")
            return
        
        if not os.environ.get("ANTHROPIC_API_KEY"):
            messagebox.showerror("API Key Missing",
                               "Please set your Anthropic API key first.")
            self._show_api_key_dialog()
            return
        
        # Disable button during review
        self.review_btn.config(state="disabled")
        self.progress_var.set(0)
        self.results_text.delete(1.0, tk.END)
        
        # Start review in separate thread
        self.review_thread = threading.Thread(target=self._run_review)
        self.review_thread.daemon = True
        self.review_thread.start()
    
    def _run_review(self):
        """Run the actual review process"""
        try:
            # Update config with selected file
            self._update_config_file()
            
            # Send status update
            self.message_queue.put(("status", "Initializing reviewer..."))
            self.message_queue.put(("progress", 10))
            
            # Initialize reviewer
            config_manager = ConfigManager(self.config_path)
            reviewer = DocumentReviewer(config_manager)
            
            # Monkey patch the reviewer to send progress updates
            original_chunk_method = reviewer._chunk_text
            def chunk_with_progress(*args, **kwargs):
                self.message_queue.put(("status", "Splitting document into chunks..."))
                self.message_queue.put(("progress", 20))
                return original_chunk_method(*args, **kwargs)
            reviewer._chunk_text = chunk_with_progress
            
            # Monitor chunk processing
            self._setup_progress_monitoring(reviewer)
            
            # Run review
            self.message_queue.put(("status", "Starting review..."))
            results = reviewer.review_document()
            
            # Process results
            self.message_queue.put(("progress", 100))
            self.message_queue.put(("status", "Review complete!"))
            self.message_queue.put(("results", results))
            
        except Exception as e:
            self.message_queue.put(("error", str(e)))
    
    def _setup_progress_monitoring(self, reviewer):
        """Setup progress monitoring for chunk processing"""
        # This is a simplified version - in production you'd want more granular progress
        original_review = reviewer.llm_client.review_text_content
        self.chunks_processed = 0
        self.total_chunks = 0
        
        def review_with_progress(*args, **kwargs):
            result = original_review(*args, **kwargs)
            self.chunks_processed += 1
            if self.total_chunks > 0:
                progress = 20 + (60 * self.chunks_processed / self.total_chunks)
                self.message_queue.put(("progress", progress))
                self.message_queue.put(("status", 
                    f"Processing chunk {self.chunks_processed}/{self.total_chunks}..."))
            return result
        
        reviewer.llm_client.review_text_content = review_with_progress
    
    def _update_config_file(self):
        """Update config file with selected document"""
        import yaml
        
        with open(self.config_path, 'r') as f:
            config = yaml.safe_load(f)
        
        config['document']['path'] = self.selected_file.get()
        
        # Set output path
        input_path = Path(self.selected_file.get())
        output_name = f"{input_path.stem}_review_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
        config['output']['path'] = str(input_path.parent / output_name)
        
        with open(self.config_path, 'w') as f:
            yaml.dump(config, f)
    
    def _process_queue(self):
        """Process messages from the review thread"""
        try:
            while True:
                msg_type, msg_data = self.message_queue.get_nowait()
                
                if msg_type == "status":
                    self.status_var.set(msg_data)
                elif msg_type == "progress":
                    self.progress_var.set(msg_data)
                elif msg_type == "results":
                    self._display_results(msg_data)
                elif msg_type == "error":
                    messagebox.showerror("Review Error", f"An error occurred:\n\n{msg_data}")
                    self.review_btn.config(state="normal")
                    self.status_var.set("Error occurred")
                    self.progress_var.set(0)
                
        except queue.Empty:
            pass
        
        # Schedule next check
        self.root.after(100, self._process_queue)
    
    def _display_results(self, results):
        """Display review results in a user-friendly format"""
        self.results_text.delete(1.0, tk.END)
        
        # Store suggestions for export
        self.suggestions = results.get('pass1_detailed_review', {}).get('suggestions', [])
        
        # Display summary
        self.results_text.insert(tk.END, "Review Summary\n", "heading")
        self.results_text.insert(tk.END, "="*50 + "\n\n")
        
        total_suggestions = len(self.suggestions)
        self.results_text.insert(tk.END, f"Total suggestions: {total_suggestions}\n\n")
        
        if total_suggestions == 0:
            self.results_text.insert(tk.END, "No suggestions found. Your document looks great!\n")
        else:
            # Group suggestions by category
            categories = {}
            for sug in self.suggestions:
                # Simple categorization based on content
                category = self._categorize_suggestion(sug)
                if category not in categories:
                    categories[category] = []
                categories[category].append(sug)
            
            # Display by category
            for category, items in sorted(categories.items()):
                self.results_text.insert(tk.END, f"\n{category} ({len(items)} items)\n", 
                                       "heading")
                self.results_text.insert(tk.END, "-"*40 + "\n\n")
                
                for i, sug in enumerate(items[:10], 1):  # Show first 10
                    self.results_text.insert(tk.END, f"{i}. ")
                    self.results_text.insert(tk.END, "Original: ", "category")
                    self.results_text.insert(tk.END, f"{sug['original']}\n", "original")
                    self.results_text.insert(tk.END, "   ")
                    self.results_text.insert(tk.END, "Suggested: ", "category")
                    self.results_text.insert(tk.END, f"{sug['improved']}\n\n", "improved")
                
                if len(items) > 10:
                    self.results_text.insert(tk.END, 
                        f"   ... and {len(items) - 10} more {category.lower()} suggestions\n\n")
        
        # Add global review if available
        global_review = results.get('pass2_global_review', {}).get('report')
        if global_review:
            self.results_text.insert(tk.END, "\n\nGlobal Document Review\n", "heading")
            self.results_text.insert(tk.END, "="*50 + "\n\n")
            self.results_text.insert(tk.END, global_review[:1000])  # First 1000 chars
            if len(global_review) > 1000:
                self.results_text.insert(tk.END, "\n\n[See full report in exported file]")
        
        # Show export buttons
        self.export_frame.grid(row=6, column=0, columnspan=3, pady=10)
        
        # Re-enable review button
        self.review_btn.config(state="normal")
        
        # Scroll to top
        self.results_text.see("1.0")
    
    def _categorize_suggestion(self, suggestion):
        """Simple categorization of suggestions"""
        text = f"{suggestion.get('original', '')} {suggestion.get('improved', '')}".lower()
        
        if any(word in text for word in ['spelling', 'spell', 'typo']):
            return "Spelling"
        elif any(word in text for word in ['grammar', 'tense', 'verb', 'noun']):
            return "Grammar"
        elif any(word in text for word in ['punctuation', 'comma', 'period', 'semicolon']):
            return "Punctuation"
        elif any(word in text for word in ['capitalization', 'capital', 'uppercase']):
            return "Capitalization"
        elif any(word in text for word in ['clarity', 'clear', 'confusing', 'awkward']):
            return "Clarity"
        elif any(word in text for word in ['consistency', 'consistent']):
            return "Consistency"
        else:
            return "Style"
    
    def _export_results(self, format):
        """Export results in various formats"""
        if not self.suggestions:
            messagebox.showinfo("No Results", "No suggestions to export.")
            return
        
        # Get save location
        base_name = Path(self.selected_file.get()).stem
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        
        if format == "docx":
            filename = filedialog.asksaveasfilename(
                defaultextension=".docx",
                filetypes=[("Word Document", "*.docx")],
                initialfile=f"{base_name}_review_{timestamp}.docx"
            )
            if filename:
                self._export_to_word(filename)
        
        elif format == "xlsx":
            filename = filedialog.asksaveasfilename(
                defaultextension=".xlsx",
                filetypes=[("Excel Workbook", "*.xlsx")],
                initialfile=f"{base_name}_review_{timestamp}.xlsx"
            )
            if filename:
                self._export_to_excel(filename)
        
        elif format == "md":
            filename = filedialog.asksaveasfilename(
                defaultextension=".md",
                filetypes=[("Markdown", "*.md")],
                initialfile=f"{base_name}_review_{timestamp}.md"
            )
            if filename:
                self._export_to_markdown(filename)
    
    def _export_to_word(self, filename):
        """Export suggestions to Word document"""
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor
            
            doc = Document()
            doc.add_heading('Document Review Report', 0)
            doc.add_paragraph(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
            doc.add_paragraph(f'Document: {Path(self.selected_file.get()).name}')
            
            doc.add_heading('Summary', level=1)
            doc.add_paragraph(f'Total suggestions: {len(self.suggestions)}')
            
            # Group by category
            categories = {}
            for sug in self.suggestions:
                cat = self._categorize_suggestion(sug)
                if cat not in categories:
                    categories[cat] = []
                categories[cat].append(sug)
            
            for category, items in sorted(categories.items()):
                doc.add_heading(f'{category} ({len(items)} items)', level=2)
                
                for i, sug in enumerate(items, 1):
                    p = doc.add_paragraph()
                    p.add_run(f'{i}. Original: ').bold = True
                    run = p.add_run(sug['original'])
                    run.font.color.rgb = RGBColor(255, 0, 0)
                    
                    p = doc.add_paragraph()
                    p.add_run('   Suggested: ').bold = True
                    run = p.add_run(sug['improved'])
                    run.font.color.rgb = RGBColor(0, 128, 0)
                    
                    doc.add_paragraph()  # Empty line
            
            doc.save(filename)
            messagebox.showinfo("Export Complete", f"Results exported to:\n{filename}")
            
        except ImportError:
            messagebox.showerror("Export Error", 
                "python-docx is required for Word export.\n"
                "Install with: pip install python-docx")
        except Exception as e:
            messagebox.showerror("Export Error", f"Failed to export: {str(e)}")
    
    def _export_to_excel(self, filename):
        """Export suggestions to Excel"""
        try:
            import pandas as pd
            
            # Prepare data
            data = []
            for sug in self.suggestions:
                data.append({
                    'Category': self._categorize_suggestion(sug),
                    'Original': sug['original'],
                    'Suggested': sug['improved'],
                    'Chunk': sug.get('chunk_index', 'N/A')
                })
            
            # Create DataFrame and save
            df = pd.DataFrame(data)
            
            with pd.ExcelWriter(filename, engine='openpyxl') as writer:
                df.to_excel(writer, sheet_name='Suggestions', index=False)
                
                # Add summary sheet
                summary_data = {
                    'Category': df['Category'].value_counts().index.tolist(),
                    'Count': df['Category'].value_counts().values.tolist()
                }
                summary_df = pd.DataFrame(summary_data)
                summary_df.to_excel(writer, sheet_name='Summary', index=False)
            
            messagebox.showinfo("Export Complete", f"Results exported to:\n{filename}")
            
        except ImportError:
            messagebox.showerror("Export Error",
                "pandas and openpyxl are required for Excel export.\n"
                "Install with: pip install pandas openpyxl")
        except Exception as e:
            messagebox.showerror("Export Error", f"Failed to export: {str(e)}")
    
    def _export_to_markdown(self, filename):
        """Export suggestions to Markdown"""
        try:
            with open(filename, 'w', encoding='utf-8') as f:
                f.write("# Document Review Report\n\n")
                f.write(f"**Generated:** {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n")
                f.write(f"**Document:** {Path(self.selected_file.get()).name}\n\n")
                
                f.write("## Summary\n\n")
                f.write(f"Total suggestions: {len(self.suggestions)}\n\n")
                
                # Group by category
                categories = {}
                for sug in self.suggestions:
                    cat = self._categorize_suggestion(sug)
                    if cat not in categories:
                        categories[cat] = []
                    categories[cat].append(sug)
                
                f.write("## Suggestions by Category\n\n")
                
                for category, items in sorted(categories.items()):
                    f.write(f"### {category} ({len(items)} items)\n\n")
                    
                    for i, sug in enumerate(items, 1):
                        f.write(f"{i}. **Original:** {sug['original']}\n")
                        f.write(f"   **Suggested:** {sug['improved']}\n\n")
            
            messagebox.showinfo("Export Complete", f"Results exported to:\n{filename}")
            
        except Exception as e:
            messagebox.showerror("Export Error", f"Failed to export: {str(e)}")


def main():
    """Main entry point"""
    # Create the main window
    root = tk.Tk()
    
    # Apply a modern style
    style = ttk.Style()
    style.theme_use('clam')  # or 'vista' on Windows
    
    # Create and run the GUI
    app = ReviewerGUI(root)
    root.mainloop()


if __name__ == "__main__":
    main()